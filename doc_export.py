"""
doc_export.py

Converts the plain-text outputs of cover_letter_engine.py and
resume_tailor_engine.py into downloadable PDF and DOCX files.

Neither generator changes -- they still just return plain text. This
module is a presentation layer on top of that text, kept deliberately
separate so the Gemini prompts stay untouched.

Public functions:
    tailored_resume_to_pdf(text)  -> bytes
    tailored_resume_to_docx(text) -> bytes
    cover_letter_to_pdf(text)     -> bytes
    cover_letter_to_docx(text)    -> bytes
"""

from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    ListFlowable,
    ListItem,
)


# ---------------------------------------------------------------------------
# Shared parsing: turns the tailored resume's plain text into typed blocks.
#
# resume_tailor_engine.py's prompt produces a consistent shape: name, then
# contact lines, then ALL-CAPS section headings (SUMMARY, EDUCATION,
# PROJECTS, ...), then paragraphs or "- " bulleted lines under each. This
# parser depends on that shape holding -- if Gemini's phrasing ever drifts,
# unrecognized lines fall back to plain paragraphs rather than raising, so
# a format drift degrades the output instead of crashing the download.
# ---------------------------------------------------------------------------
def _is_heading_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 40:
        return False
    if not any(c.isalpha() for c in stripped):
        return False
    return stripped == stripped.upper()


def _parse_resume_blocks(text: str):
    lines = text.strip().split("\n")
    blocks = []
    seen_heading = False
    preamble = []  # name + contact lines, collected before the first heading

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if _is_heading_line(line):
            if not seen_heading and preamble:
                blocks.append(("name", preamble[0]))
                for c in preamble[1:]:
                    blocks.append(("contact", c))
                preamble = []
            seen_heading = True
            blocks.append(("heading", line))
            continue

        if not seen_heading:
            preamble.append(line)
            continue

        if line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        else:
            blocks.append(("paragraph", line))

    # No ALL-CAPS heading detected anywhere -- treat the whole thing as
    # name + plain paragraphs rather than losing the content.
    if not seen_heading and preamble:
        blocks.append(("name", preamble[0]))
        for c in preamble[1:]:
            blocks.append(("paragraph", c))

    return blocks


import re

_CLOSING_PATTERN = re.compile(
    r"^(sincerely|regards|best regards|warm regards|best|thank you|thanks|"
    r"yours truly|yours sincerely|respectfully)[,]?$",
    re.IGNORECASE,
)


def _paragraphs(text: str):
    """
    Split prose (cover letters) into paragraph blocks on blank lines.

    Each block is returned as a list of lines. In most blocks, those
    lines are soft-wrapped and should be joined into one flowing line
    of text. The exception is a signature block ("Sincerely," on its
    own line followed by a name) -- collapsing that into one line reads
    wrong, so closing-salutation blocks keep their line breaks intact.
    """
    raw_blocks = [b.strip() for b in text.strip().split("\n\n") if b.strip()]
    result = []
    for block in raw_blocks:
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if lines and _CLOSING_PATTERN.match(lines[0]):
            result.append(lines)  # keep as separate lines
        else:
            result.append([" ".join(lines)])  # collapse into one flowing line
    return result


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------
def tailored_resume_to_docx(text: str) -> bytes:
    blocks = _parse_resume_blocks(text)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for kind, content in blocks:
        if kind == "name":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
        elif kind == "contact":
            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.size = Pt(9.5)
        elif kind == "heading":
            doc.add_heading(content.title(), level=2)
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        elif kind == "paragraph":
            doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def cover_letter_to_docx(text: str) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for block_lines in _paragraphs(text):
        p = doc.add_paragraph()
        for i, line in enumerate(block_lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------
def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ResumeName", fontSize=18, leading=22, spaceAfter=4,
        fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="ResumeContact", fontSize=9.5, leading=13,
        textColor=colors.HexColor("#444444"), spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name="ResumeHeading", fontSize=12, leading=16, spaceBefore=12, spaceAfter=4,
        fontName="Helvetica-Bold", textColor=colors.HexColor("#7c6a3a"),
    ))
    styles.add(ParagraphStyle(
        name="ResumeBody", fontSize=10, leading=14, spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="LetterBody", fontSize=11, leading=16, spaceAfter=12,
    ))
    return styles


def tailored_resume_to_pdf(text: str) -> bytes:
    blocks = _parse_resume_blocks(text)
    styles = _pdf_styles()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
    )

    story = []
    bullet_group = []

    def flush_bullets():
        nonlocal bullet_group
        if bullet_group:
            story.append(ListFlowable(
                [ListItem(Paragraph(b, styles["ResumeBody"])) for b in bullet_group],
                bulletType="bullet", leftIndent=14,
            ))
            bullet_group = []

    for kind, content in blocks:
        if kind == "bullet":
            bullet_group.append(content)
            continue
        flush_bullets()

        if kind == "name":
            story.append(Paragraph(content, styles["ResumeName"]))
        elif kind == "contact":
            story.append(Paragraph(content, styles["ResumeContact"]))
        elif kind == "heading":
            story.append(Paragraph(content.title(), styles["ResumeHeading"]))
        elif kind == "paragraph":
            story.append(Paragraph(content, styles["ResumeBody"]))

    flush_bullets()
    doc.build(story)
    return buffer.getvalue()


def cover_letter_to_pdf(text: str) -> bytes:
    styles = _pdf_styles()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=1 * inch, rightMargin=1 * inch,
        topMargin=1 * inch, bottomMargin=1 * inch,
    )

    story = [
        Paragraph("<br/>".join(block_lines), styles["LetterBody"])
        for block_lines in _paragraphs(text)
    ]
    doc.build(story)
    return buffer.getvalue()